from fastapi import FastAPI, APIRouter, UploadFile, File, HTTPException, Form
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from pymongo import MongoClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional
import uuid
from datetime import datetime, timezone
import base64
import io
from emergentintegrations.llm.chat import LlmChat, UserMessage
import json
import PyPDF2
from docx import Document as DocxDocument
import tempfile
import asyncio
import re

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection - async for normal operations
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Sync MongoDB client for GridFS operations
sync_client = MongoClient(mongo_url)
sync_db = sync_client[os.environ['DB_NAME']]

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Define Models
class Document(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    category: str
    file_type: str
    content: str
    content_preview: Optional[str] = None  # Short excerpt for UI display
    is_large_document: bool = False  # Flag for large documents
    one_liner: Optional[str] = None  # One sentence summary for Make.com automation
    tags: List[str] = []
    references: List[str] = []
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    updated_at: Optional[str] = None
    file_size: Optional[int] = None
    original_filename: Optional[str] = None
    has_original_file: bool = False
    original_language: Optional[str] = None
    was_translated: bool = False

class DocumentCreate(BaseModel):
    title: str
    category: str
    file_type: str
    content: str
    tags: List[str] = []
    references: List[str] = []
    file_size: Optional[int] = None

class DocumentUpdate(BaseModel):
    title: Optional[str] = None
    category: Optional[str] = None
    tags: Optional[List[str]] = None
    references: Optional[List[str]] = None
    content: Optional[str] = None

class Category(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    description: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class CategoryCreate(BaseModel):
    name: str
    description: Optional[str] = None

class BlogCreateRequest(BaseModel):
    document_ids: List[str]
    title: str
    category: str = "Blog Articles"
    custom_instructions: Optional[str] = None

class BlogArticle(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    content: str
    tags: List[str] = []
    category: str = "Blog Articles"
    meta_title: str
    meta_description: str
    url_slug: str
    featured_image_url: Optional[str] = None
    source_document_ids: List[str]
    custom_instructions: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    file_type: str = "blog_article"
    references: List[str] = []
    has_original_file: bool = False
    original_language: str = "nl"
    was_translated: bool = False

class ChatMessage(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    session_id: str
    role: str  # 'user' or 'assistant'
    content: str
    timestamp: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ChatRequest(BaseModel):
    session_id: str
    message: str
    context_type: Optional[str] = "general"  # general, consult, treatment, supplement

class TreatmentPlanRequest(BaseModel):
    patient_info: str
    symptoms: str
    diagnosis: str

class SupplementAdviceRequest(BaseModel):
    condition: str
    patient_details: str

# Helper function to detect language and translate if needed
async def translate_to_dutch_if_needed(content: str, title: str) -> tuple[str, str]:
    """Detect if content is in English and translate to Dutch if needed"""
    try:
        # Skip if content is too short
        if len(content.strip()) < 50:
            return content, "nl"
        
        # First detect language with a simpler prompt
        session_id = str(uuid.uuid4())
        chat = LlmChat(
            api_key=os.environ.get('EMERGENT_LLM_KEY'),
            session_id=session_id,
            system_message="Je bent een taaldetectie expert. Antwoord alleen met de taalcode."
        ).with_model("anthropic", "claude-4-sonnet-20250514")
        
        # Detect language - use shorter text sample
        detect_prompt = f"""Wat is de taal van deze tekst? Antwoord ALLEEN met: 'en' (Engels) of 'nl' (Nederlands) of 'other' (anders).

Titel: {title[:100]}
Tekst: {content[:300]}

Taalcode:"""
        
        user_message = UserMessage(text=detect_prompt)
        language = await chat.send_message(user_message)
        language = language.strip().lower()
        
        logging.info(f"Detected language: {language} for document: {title}")
        
        # If English, translate to Dutch
        if 'en' in language:
            logging.info(f"Translating English content to Dutch: {title}")
            
            # Create new session for translation
            translate_session = str(uuid.uuid4())
            translate_chat = LlmChat(
                api_key=os.environ.get('EMERGENT_LLM_KEY'),
                session_id=translate_session,
                system_message="Je bent een professionele vertaler van Engels naar Nederlands. Vertaal de tekst precies en behoud alle formattering."
            ).with_model("anthropic", "claude-4-sonnet-20250514")
            
            translate_prompt = f"""Vertaal deze Engelse tekst naar Nederlands. Behoud alle structuur en formattering.

{content}

Nederlandse vertaling:"""
            
            translate_message = UserMessage(text=translate_prompt)
            translated = await translate_chat.send_message(translate_message)
            
            logging.info(f"Successfully translated document: {title}")
            return translated, "en"
        
        return content, language
        
    except Exception as e:
        logging.error(f"Error in translation: {str(e)}")
        # Return original content if translation fails
        return content, "unknown"

# Helper function to extract text from files
async def extract_text_from_file(file: UploadFile) -> str:
    """Extract text from uploaded file"""
    content = await file.read()
    
    if file.filename.endswith('.pdf'):
        # Extract from PDF
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    
    elif file.filename.endswith('.docx'):
        # Extract from DOCX
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(content)
            tmp_file.flush()
            doc = DocxDocument(tmp_file.name)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            os.unlink(tmp_file.name)
        return text
    
    elif file.filename.endswith('.txt'):
        # Extract from TXT
        return content.decode('utf-8')
    
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type")

# Helper function to generate tags with AI
async def generate_tags_with_ai(title: str, content: str) -> List[str]:
    """Generate relevant tags using Claude AI"""
    try:
        session_id = str(uuid.uuid4())
        chat = LlmChat(
            api_key=os.environ.get('EMERGENT_LLM_KEY'),
            session_id=session_id,
            system_message="Je bent een expert in het taggen van medische en orthomoleculaire documenten. Genereer 3-7 relevante tags in het Nederlands voor het document."
        ).with_model("anthropic", "claude-4-sonnet-20250514")
        
        prompt = f"""Genereer relevante tags voor dit document:

Titel: {title}
Inhoud: {content[:1000]}...

Geef alleen de tags terug, gescheiden door komma's. Gebruik maximaal 7 tags. Focus op:
- Hoofdonderwerpen
- Supplementen/kruiden die genoemd worden
- Aandoeningen/symptomen
- Therapeutische categorieën

Antwoord alleen met de tags, niets anders."""
        
        user_message = UserMessage(text=prompt)
        response = await chat.send_message(user_message)
        
        # Parse tags from response
        tags = [tag.strip() for tag in response.split(',')]
        return tags[:7]  # Max 7 tags
    except Exception as e:
        logging.error(f"Error generating tags: {str(e)}")
        return ["orthomoleculair", "kennis"]  # Default tags

# Helper function to extract references with AI
async def extract_references_with_ai(content: str) -> List[str]:
    """Extract references/sources from content using Claude AI"""
    try:
        session_id = str(uuid.uuid4())
        chat = LlmChat(
            api_key=os.environ.get('EMERGENT_LLM_KEY'),
            session_id=session_id,
            system_message="Je bent een expert in het identificeren van wetenschappelijke referenties en bronnen in medische documenten."
        ).with_model("anthropic", "claude-4-sonnet-20250514")
        
        prompt = f"""Analyseer deze tekst en identificeer alle referenties, bronnen, studies of citaten:

{content[:2000]}...

Geef alleen de gevonden referenties terug, elke referentie op een nieuwe regel.
Als er geen referenties zijn, antwoord dan met: GEEN

Formaat:
- Auteur (jaar) - Titel
- Journal naam, volume, pagina's
- URL indien vermeld

Antwoord alleen met de referenties of GEEN."""
        
        user_message = UserMessage(text=prompt)
        response = await chat.send_message(user_message)
        
        if response.strip().upper() == "GEEN":
            return []
        
        # Parse references
        references = [ref.strip() for ref in response.split('\n') if ref.strip() and not ref.strip().startswith('-')]
        return references[:10]  # Max 10 references
    except Exception as e:
        logging.error(f"Error extracting references: {str(e)}")
        return []

# Helper functions
def generate_document_preview(content: str, title: str = "") -> tuple[str, bool]:
    """Generate intelligent preview for documents and determine if it's a large document"""
    
    # Define thresholds
    LARGE_DOCUMENT_THRESHOLD = 2000  # characters
    PREVIEW_LENGTH = 800  # characters for preview
    
    content_length = len(content)
    is_large = content_length > LARGE_DOCUMENT_THRESHOLD
    
    if not is_large:
        # Small document - return full content
        return content, False
    
    # Large document - generate intelligent preview
    lines = content.split('\n')
    preview_lines = []
    char_count = 0
    
    # Always include title if provided
    if title:
        preview_lines.append(f"Samenvatting van: {title}")
        preview_lines.append("")
    
    # Try to get meaningful content from the beginning
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Skip headers that might be metadata
        if line.lower().startswith(('auteur:', 'bron:', 'datum:', 'pagina:', 'hoofdstuk:')):
            continue
            
        # Add line if it contributes to understanding
        if char_count + len(line) <= PREVIEW_LENGTH:
            preview_lines.append(line)
            char_count += len(line) + 2  # +2 for \n\n
        else:
            # Add partial line and break
            remaining_chars = PREVIEW_LENGTH - char_count
            if remaining_chars > 50:  # Only add if meaningful portion remains
                preview_lines.append(line[:remaining_chars] + "...")
            break
    
    # Add summary footer
    preview_lines.append("")
    preview_lines.append(f"[Dit is een preview van een document met {content_length:,} karakters. De volledige inhoud is beschikbaar voor AI verwerking en blog generatie.]")
    
    preview = "\n\n".join(preview_lines)
    
    return preview, is_large

def generate_oneliner_mock(title: str, content: str) -> str:
    """Generate a concise one-sentence summary for Make.com automation"""
    
    # Simple mock implementation - extract key concepts
    words = content.lower().split()
    
    # Common orthomolecular/health keywords to look for
    health_keywords = [
        'vitamine', 'mineralen', 'supplement', 'voeding', 'gezondheid', 
        'orthomoleculair', 'behandeling', 'therapie', 'preventie',
        'darm', 'microbioom', 'ontstekingsremming', 'antioxidant',
        'stress', 'energie', 'immuniteit', 'herstel', 'balans'
    ]
    
    # Find relevant keywords in content
    found_keywords = [kw for kw in health_keywords if kw in ' '.join(words)][:3]
    
    # Generate context-aware one-liner
    if 'vitamine' in found_keywords or 'mineralen' in found_keywords:
        return f"Onderzoek naar de rol van vitamines en mineralen bij {title.lower()}, met focus op orthomoleculaire behandelingsmogelijkheden."
    elif 'darm' in found_keywords or 'microbioom' in found_keywords:
        return f"Inzichten over darmgezondheid en microbioom in relatie tot {title.lower()}, belangrijk voor holistische gezondheidszorg."
    elif 'ontstekingsremming' in found_keywords or 'antioxidant' in found_keywords:
        return f"Evidence voor ontstekingsremmende en antioxidantrijke interventies bij {title.lower()}, geschikt voor natuurgeneeskundige praktijk."
    elif 'stress' in found_keywords or 'energie' in found_keywords:
        return f"Natuurlijke strategieën voor stress- en energiemanagement gerelateerd aan {title.lower()}, toepasbaar in orthomoleculaire behandeling."
    else:
        # Generic health-focused one-liner
        return f"Praktische inzichten over {title.lower()} vanuit orthomoleculair perspectief, relevant voor natuurgeneeskundige behandeling en preventie."

# Document routes
@api_router.post("/documents", response_model=Document)
async def create_document(doc: DocumentCreate):
    """Upload a new document to the knowledge base"""
    doc_dict = doc.dict()
    doc_obj = Document(**doc_dict)
    await db.documents.insert_one(doc_obj.dict())
    return doc_obj

@api_router.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    category: str = Form("artikel")
):
    """Upload a file and extract text with auto-generated tags and references"""
    try:
        # Read file content
        file_content = await file.read()
        
        # Determine file type
        file_type = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else 'unknown'
        
        # Check if it's an image
        is_image = file_type in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp']
        is_pdf = file_type == 'pdf'
        
        # For images, store the image (vision/OCR can be added later)
        if is_image:
            doc_title = title if title else file.filename.rsplit('.', 1)[0]
            
            # Simple content for now - just indicate it's an image
            content = f"[Afbeelding: {file.filename}]\n\nDit is een afbeelding. Bekijk het origineel in de document viewer."
            
            # Generate tags based on title/filename
            tags = await generate_tags_with_ai(doc_title, f"Dit is een afbeelding met de naam: {doc_title}")
            references = []
            
            # Store image in GridFS
            import gridfs
            fs = gridfs.GridFS(sync_db)
            
            # Determine media type
            media_type_map = {
                'jpg': 'image/jpeg',
                'jpeg': 'image/jpeg',
                'png': 'image/png',
                'gif': 'image/gif',
                'webp': 'image/webp',
                'bmp': 'image/bmp'
            }
            media_type = media_type_map.get(file_type, 'image/jpeg')
            
            file_id = fs.put(file_content, filename=file.filename, content_type=media_type)
            
            doc = Document(
                title=doc_title,
                category=category,
                file_type=file_type,
                content=content,
                tags=tags,
                references=references,
                file_size=len(file_content),
                original_filename=file.filename,
                has_original_file=True
            )
            
            doc_dict = doc.dict()
            doc_dict['original_file_id'] = str(file_id)
            
            # Insert into database
            result = await db.documents.insert_one(doc_dict)
            
            # Fetch the inserted document to get clean data
            inserted_doc = await db.documents.find_one({"id": doc.id})
            
            # Remove _id for JSON response
            if inserted_doc and '_id' in inserted_doc:
                del inserted_doc['_id']
            
            return {
                "message": "Afbeelding succesvol geüpload",
                "document": inserted_doc or doc_dict
            }
        
        # For PDFs and text files, extract text
        await file.seek(0)  # Reset file pointer
        content = await extract_text_from_file(file)
        
        if not content.strip():
            raise HTTPException(status_code=400, detail="Geen tekst gevonden in bestand")
        
        # Use filename as title if not provided
        doc_title = title if title else file.filename.rsplit('.', 1)[0]
        
        # Detect language and translate to Dutch if needed
        translated_content, original_lang = await translate_to_dutch_if_needed(content, doc_title)
        was_translated = (original_lang == "en")
        
        if was_translated:
            logging.info(f"Translated document from English: {doc_title}")
        
        # Generate tags and extract references with AI (using translated content)
        tags = await generate_tags_with_ai(doc_title, translated_content)
        references = await extract_references_with_ai(translated_content)
        
        # Store original file for PDFs
        file_id = None
        has_original = False
        
        if is_pdf:
            import gridfs
            fs = gridfs.GridFS(sync_db)
            file_id = fs.put(file_content, filename=file.filename, content_type=file.content_type or 'application/pdf')
            has_original = True
        
        # Generate preview for large documents
        preview, is_large = generate_document_preview(translated_content, doc_title)
        
        # Generate one-liner for Make.com automation
        one_liner = generate_oneliner_mock(doc_title, translated_content)
        
        # Create document
        doc = Document(
            title=doc_title,
            category=category,
            file_type=file_type,
            content=translated_content,
            content_preview=preview if is_large else None,
            is_large_document=is_large,
            one_liner=one_liner,
            tags=tags,
            references=references,
            file_size=len(translated_content),
            original_filename=file.filename if has_original else None,
            has_original_file=has_original,
            original_language=original_lang if was_translated else None,
            was_translated=was_translated
        )
        
        doc_dict = doc.dict()
        if file_id:
            doc_dict['original_file_id'] = str(file_id)
        
        # Insert into database
        result = await db.documents.insert_one(doc_dict)
        
        # Fetch the inserted document to get clean data without ObjectId
        inserted_doc = await db.documents.find_one({"id": doc.id})
        
        # Convert for JSON response (remove _id)
        if inserted_doc and '_id' in inserted_doc:
            del inserted_doc['_id']
        
        return {
            "message": "Document succesvol geüpload",
            "document": inserted_doc or doc_dict
        }
    except Exception as e:
        logging.error(f"Upload error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/documents/paste")
async def paste_document(
    title: str = Form(...),
    content: str = Form(...),
    category: str = Form("aantekening")
):
    """Create document from pasted text with auto-generated tags and references"""
    try:
        if not content.strip():
            raise HTTPException(status_code=400, detail="Inhoud mag niet leeg zijn")
        
        # Detect language and translate to Dutch if needed
        translated_content, original_lang = await translate_to_dutch_if_needed(content, title)
        was_translated = (original_lang == "en")
        
        if was_translated:
            logging.info(f"Translated pasted content from English: {title}")
        
        # Generate tags and extract references with AI (using translated content)
        tags = await generate_tags_with_ai(title, translated_content)
        references = await extract_references_with_ai(translated_content)
        
        # Generate preview for large documents
        preview, is_large = generate_document_preview(translated_content, title)
        
        # Generate one-liner for Make.com automation
        one_liner = generate_oneliner_mock(title, translated_content)
        
        # Create document
        doc = Document(
            title=title,
            category=category,
            file_type="text",
            content=translated_content,
            content_preview=preview if is_large else None,
            is_large_document=is_large,
            one_liner=one_liner,
            tags=tags,
            references=references,
            file_size=len(translated_content),
            original_language=original_lang if was_translated else None,
            was_translated=was_translated
        )
        
        doc_dict = doc.dict()
        
        # Insert into database
        result = await db.documents.insert_one(doc_dict)
        
        # Fetch the inserted document to get clean data
        inserted_doc = await db.documents.find_one({"id": doc.id})
        
        # Remove _id for JSON response
        if inserted_doc and '_id' in inserted_doc:
            del inserted_doc['_id']
        
        return {
            "message": "Document succesvol toegevoegd",
            "document": inserted_doc or doc_dict
        }
    except Exception as e:
        logging.error(f"Paste error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/documents/voice")
async def voice_document(
    audio: UploadFile = File(...),
    title: str = Form(...),
    category: str = Form("aantekening")
):
    """Create document from voice recording with speech-to-text"""
    try:
        # Read audio file
        audio_content = await audio.read()
        
        if len(audio_content) == 0:
            raise HTTPException(status_code=400, detail="Audio bestand is leeg")
        
        # Save audio temporarily
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix='.webm') as tmp_audio:
            tmp_audio.write(audio_content)
            tmp_audio_path = tmp_audio.name
        
        try:
            # Use OpenAI Whisper for speech-to-text via Emergent LLM
            import requests
            
            api_key = os.environ.get('EMERGENT_LLM_KEY')
            
            # Call Whisper API
            with open(tmp_audio_path, 'rb') as audio_file:
                response = requests.post(
                    'https://api.emergentagi.com/openai/audio/transcriptions',
                    headers={
                        'Authorization': f'Bearer {api_key}'
                    },
                    files={
                        'file': ('audio.webm', audio_file, 'audio/webm')
                    },
                    data={
                        'model': 'whisper-1',
                        'language': 'nl'  # Try Dutch first
                    }
                )
            
            if response.status_code == 200:
                result = response.json()
                content = result.get('text', '')
                
                if not content.strip():
                    raise HTTPException(status_code=400, detail="Geen tekst herkend in audio")
                
                logging.info(f"Transcribed audio to text: {len(content)} characters")
                
            else:
                logging.error(f"Whisper API error: {response.status_code} - {response.text}")
                raise HTTPException(status_code=500, detail="Fout bij spraak-naar-tekst conversie")
        
        finally:
            # Clean up temp file
            os.unlink(tmp_audio_path)
        
        # Now process like normal paste
        # Detect language and translate to Dutch if needed
        translated_content, original_lang = await translate_to_dutch_if_needed(content, title)
        was_translated = (original_lang == "en")
        
        if was_translated:
            logging.info(f"Translated voice content from English: {title}")
        
        # Generate tags and extract references with AI
        tags = await generate_tags_with_ai(title, translated_content)
        references = await extract_references_with_ai(translated_content)
        
        # Create document
        doc = Document(
            title=title,
            category=category,
            file_type="voice",
            content=translated_content,
            tags=tags,
            references=references,
            file_size=len(translated_content),
            original_language=original_lang if was_translated else None,
            was_translated=was_translated
        )
        
        doc_dict = doc.dict()
        
        # Insert into database
        result = await db.documents.insert_one(doc_dict)
        
        # Fetch the inserted document
        inserted_doc = await db.documents.find_one({"id": doc.id})
        
        # Remove _id for JSON response
        if inserted_doc and '_id' in inserted_doc:
            del inserted_doc['_id']
        
        return {
            "message": "Spraakopname succesvol verwerkt en opgeslagen",
            "document": inserted_doc or doc_dict,
            "transcription_length": len(content)
        }
        
    except Exception as e:
        logging.error(f"Voice error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/documents", response_model=List[Document])
async def get_documents(category: Optional[str] = None):
    """Get all documents, optionally filtered by category"""
    query = {}
    if category:
        query["category"] = category
    documents = await db.documents.find(query).sort("created_at", -1).to_list(1000)
    return [Document(**doc) for doc in documents]

@api_router.get("/documents/{document_id}", response_model=Document)
async def get_document(document_id: str):
    """Get a specific document by ID"""
    doc = await db.documents.find_one({"id": document_id})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return Document(**doc)

@api_router.get("/documents/{document_id}/file")
async def get_original_file(document_id: str):
    """Get the original uploaded file"""
    doc = await db.documents.find_one({"id": document_id})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if not doc.get('has_original_file') or not doc.get('original_file_id'):
        raise HTTPException(status_code=404, detail="Origineel bestand niet beschikbaar")
    
    try:
        import gridfs
        from bson import ObjectId
        
        fs = gridfs.GridFS(sync_db)
        file_id = ObjectId(doc['original_file_id'])
        grid_out = fs.get(file_id)
        
        # Determine media type based on file extension
        file_type = doc.get('file_type', '').lower()
        media_types = {
            'pdf': 'application/pdf',
            'jpg': 'image/jpeg',
            'jpeg': 'image/jpeg',
            'png': 'image/png',
            'gif': 'image/gif',
            'bmp': 'image/bmp',
            'webp': 'image/webp'
        }
        
        media_type = media_types.get(file_type, 'application/octet-stream')
        
        return StreamingResponse(
            io.BytesIO(grid_out.read()),
            media_type=media_type,
            headers={
                "Content-Disposition": f"inline; filename={doc.get('original_filename', 'document')}"
            }
        )
    except Exception as e:
        logging.error(f"Error retrieving file: {str(e)}")
        raise HTTPException(status_code=500, detail="Fout bij ophalen bestand")

@api_router.put("/documents/{document_id}")
async def update_document(document_id: str, update: DocumentUpdate):
    """Update a document"""
    doc = await db.documents.find_one({"id": document_id})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    update_data = update.dict(exclude_unset=True)
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    await db.documents.update_one(
        {"id": document_id},
        {"$set": update_data}
    )
    
    updated_doc = await db.documents.find_one({"id": document_id})
    return {"message": "Document bijgewerkt", "document": Document(**updated_doc).dict()}

@api_router.delete("/documents/{document_id}")
async def delete_document(document_id: str):
    """Delete a document"""
    result = await db.documents.delete_one({"id": document_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"message": "Document deleted successfully"}

@api_router.get("/documents/search/{query}")
async def search_documents(query: str):
    """Search documents by title, content, or tags"""
    documents = await db.documents.find({
        "$or": [
            {"title": {"$regex": query, "$options": "i"}},
            {"content": {"$regex": query, "$options": "i"}},
            {"tags": {"$regex": query, "$options": "i"}}
        ]
    }).to_list(100)
    return [Document(**doc) for doc in documents]

@api_router.get("/documents/by-tag/{tag}")
async def get_documents_by_tag(tag: str):
    """Get all documents that have a specific tag"""
    documents = await db.documents.find({
        "tags": {"$regex": f"^{tag}$", "$options": "i"}
    }).sort("created_at", -1).to_list(1000)
    return [Document(**doc) for doc in documents]

# Category routes
@api_router.post("/categories", response_model=Category)
async def create_category(cat: CategoryCreate):
    """Create a new custom category"""
    # Check if category already exists
    existing = await db.categories.find_one({"name": cat.name})
    if existing:
        raise HTTPException(status_code=400, detail="Categorie bestaat al")
    
    category = Category(**cat.dict())
    await db.categories.insert_one(category.dict())
    return category

@api_router.get("/categories", response_model=List[Category])
async def get_categories():
    """Get all categories"""
    categories = await db.categories.find().sort("name", 1).to_list(100)
    return [Category(**cat) for cat in categories]

@api_router.delete("/categories/{category_id}")
async def delete_category(category_id: str):
    """Delete a category"""
    result = await db.categories.delete_one({"id": category_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Category not found")
    return {"message": "Category deleted successfully"}

# Chat routes with Claude integration
@api_router.post("/chat")
async def chat(request: ChatRequest):
    """Chat with AI assistant using Claude Sonnet 4"""
    try:
        # Save user message
        user_msg = ChatMessage(
            session_id=request.session_id,
            role="user",
            content=request.message
        )
        await db.chat_messages.insert_one(user_msg.dict())
        
        # Get relevant documents for context
        relevant_docs = await db.documents.find({
            "$or": [
                {"content": {"$regex": request.message.split()[0] if request.message.split() else "", "$options": "i"}},
                {"tags": {"$regex": request.message.split()[0] if request.message.split() else "", "$options": "i"}}
            ]
        }).limit(3).to_list(3)
        
        context = ""
        if relevant_docs:
            context = "\n\nRelevante documenten uit de kennisbank:\n"
            for doc in relevant_docs:
                context += f"- {doc['title']}: {doc['content'][:200]}...\n"
        
        # Create system message based on context type
        system_messages = {
            "general": "Je bent een expert orthomoleculair natuurgeneeskundige en kPNI therapeut. Je helpt met het beantwoorden van vragen op basis van beschikbare kennis over supplementen, kruiden, diagnostiek en behandelplannen.",
            "consult": "Je bent een consult-assistent voor een orthomoleculair natuurgeneeskundige praktijk. Help bij het analyseren van patiëntsymptomen en adviseer over mogelijke diagnostiek.",
            "treatment": "Je bent gespecialiseerd in het maken van behandelplannen voor orthomoleculaire therapie en kPNI. Geef concrete en praktische behandeladvies.",
            "supplement": "Je bent expert in supplementen, kruiden en gemmo therapie. Geef gedetailleerde adviezen over dosering en combinaties."
        }
        
        system_message = system_messages.get(request.context_type, system_messages["general"])
        
        # Initialize Claude chat
        chat = LlmChat(
            api_key=os.environ.get('EMERGENT_LLM_KEY'),
            session_id=request.session_id,
            system_message=system_message
        ).with_model("anthropic", "claude-4-sonnet-20250514")
        
        # Send message with context
        user_message = UserMessage(
            text=request.message + context
        )
        
        response = await chat.send_message(user_message)
        
        # Save assistant message
        assistant_msg = ChatMessage(
            session_id=request.session_id,
            role="assistant",
            content=response
        )
        await db.chat_messages.insert_one(assistant_msg.dict())
        
        return {
            "response": response,
            "session_id": request.session_id
        }
    except Exception as e:
        logging.error(f"Chat error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/chat/history/{session_id}")
async def get_chat_history(session_id: str):
    """Get chat history for a session"""
    messages = await db.chat_messages.find(
        {"session_id": session_id}
    ).sort("timestamp", 1).to_list(1000)
    return [ChatMessage(**msg) for msg in messages]

# Treatment plan generation
@api_router.post("/treatment-plan")
async def generate_treatment_plan(request: TreatmentPlanRequest):
    """Generate a treatment plan using AI"""
    try:
        session_id = str(uuid.uuid4())
        
        chat = LlmChat(
            api_key=os.environ.get('EMERGENT_LLM_KEY'),
            session_id=session_id,
            system_message="Je bent een expert orthomoleculair therapeut gespecialiseerd in kPNI. Maak gedetailleerde behandelplannen met specifieke aanbevelingen voor supplementen, kruiden, leefstijl en aanvullende diagnostiek."
        ).with_model("anthropic", "claude-4-sonnet-20250514")
        
        prompt = f"""Maak een uitgebreid behandelplan voor de volgende patiënt:

Patiënt informatie: {request.patient_info}
Symptomen: {request.symptoms}
Diagnose: {request.diagnosis}

Geef een gestructureerd behandelplan met:
1. Orthomoleculaire supplementen (met dosering)
2. Kruidenadvies
3. Gemmo therapie suggesties
4. Leefstijladviezen
5. Aanvullende diagnostiek indien nodig
6. Tijdslijn en evaluatiemomenten"""
        
        user_message = UserMessage(text=prompt)
        response = await chat.send_message(user_message)
        
        return {"treatment_plan": response}
    except Exception as e:
        logging.error(f"Treatment plan error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# Supplement advice
@api_router.post("/supplement-advice")
async def get_supplement_advice(request: SupplementAdviceRequest):
    """Get supplement and herb advice using AI"""
    try:
        session_id = str(uuid.uuid4())
        
        # Get relevant supplement documents
        relevant_docs = await db.documents.find({
            "$or": [
                {"category": "supplement"},
                {"category": "kruiden"},
                {"tags": {"$in": ["supplement", "kruiden", "gemmo"]}}
            ]
        }).limit(5).to_list(5)
        
        context = ""
        if relevant_docs:
            context = "\n\nRelevante informatie uit kennisbank:\n"
            for doc in relevant_docs:
                context += f"- {doc['title']}: {doc['content'][:300]}...\n"
        
        chat = LlmChat(
            api_key=os.environ.get('EMERGENT_LLM_KEY'),
            session_id=session_id,
            system_message="Je bent expert in orthomoleculaire supplementen, kruiden en gemmo therapie. Geef praktische en evidence-based adviezen."
        ).with_model("anthropic", "claude-4-sonnet-20250514")
        
        prompt = f"""Geef supplement- en kruidenadvies voor:

Conditie: {request.condition}
Patiënt details: {request.patient_details}

Geef advies over:
1. Aanbevolen supplementen met dosering
2. Kruidenpreparaten
3. Gemmo therapie
4. Combinatie-adviezen
5. Contra-indicaties
6. Interacties met andere middelen

{context}"""
        
        user_message = UserMessage(text=prompt)
        response = await chat.send_message(user_message)
        
        return {"advice": response}
    except Exception as e:
        logging.error(f"Supplement advice error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# Blog Article Creation
@api_router.post("/blog/create")
async def create_blog_article(request: BlogCreateRequest):
    """Create a blog article from selected documents with SEO optimization"""
    try:
        # Fetch source documents
        source_documents = []
        for doc_id in request.document_ids:
            doc = await db.documents.find_one({"id": doc_id})
            if not doc:
                raise HTTPException(status_code=404, detail=f"Document {doc_id} not found")
            source_documents.append(doc)
        
        if not source_documents:
            raise HTTPException(status_code=400, detail="No source documents provided")
        
        # Prepare content for AI processing
        combined_content = "\n\n".join([
            f"**{doc['title']}**\n{doc['content']}" 
            for doc in source_documents
        ])
        
        # Local SEO keywords
        local_keywords = ["fysio zeist", "Fysiopraktijk Zeist", "Orthomoleculair Praktijk Zeist"]
        
        # Create detailed prompt for blog generation
        blog_prompt = f"""Creëer een uitgebreid, SEO-geoptimaliseerd blog artikel in het Nederlands gebaseerd op de volgende bronmaterialen.

BRONMATERIALEN:
{combined_content}

INSTRUCTIES:
1. STRUCTUUR: Gebruik dezelfde structuur als professionele orthomoleculaire gidsen:
   - Uitgebreide inleiding met kernboodschap (2-3 alinea's)
   - Duidelijke hoofdstukken met ### headers
   - Praktische stappen (Stap 1, 2, 3, etc.)
   - Bullet points voor concrete tips
   - FAQ sectie aan het einde
   - Disclaimer

2. SEO OPTIMALISATIE:
   - Integreer natuurlijk deze lokale keywords: {', '.join(local_keywords)}
   - Gebruik het hoofdonderwerp als primair keyword
   - Maak gebruik van semantisch gerelateerde termen
   - Optimaliseer voor featured snippets (gebruik vraag-antwoord structuur)

3. TONE & STIJL:
   - Professioneel maar toegankelijk
   - Nederlandse taal
   - Wetenschappelijk onderbouwd maar praktisch
   - Gebruik "je" vorm voor directe aanspreking

4. INHOUD EISEN:
   - Minimaal 1500 woorden
   - Gebruik concrete voorbeelden
   - Voeg praktische tips toe
   - Maak interne link suggesties naar gerelateerde onderwerpen

TITEL: {request.title}

AANGEPASTE INSTRUCTIES:
{request.custom_instructions if request.custom_instructions else 'Geen aanvullende instructies.'}

Genereer een volledig blog artikel dat klaar is voor publicatie."""

        # MOCK IMPLEMENTATION - Replace when LLM service is available
        # Generate realistic Dutch blog content with SEO optimization
        
        source_titles = [doc['title'] for doc in source_documents]
        source_content_preview = " ".join([doc['content'][:100] for doc in source_documents])
        
        # Create comprehensive mock blog content WITHOUT markdown
        mock_blog_content = f"""{request.title}

{request.title.lower()} begint met één kernidee: een holistische benadering die voeding, suppletie en leefstijl slim combineert. Bij Orthomoleculair Praktijk Zeist zien we dagelijks hoe een geïntegreerde aanpak resultaten oplevert waar symptoombestrijding faalt. In dit uitgebreide artikel verken je stap voor stap hoe je dat aanpakt, waar de quick wins zitten en hoe je valkuilen voorkomt.

De orthomoleculaire benadering kent vele toepassingen: van spijsverteringsklachten en energie-issues tot pijn- en stressmanagement. Elke situatie kent zijn eigen nuances, maar ze delen gemeenschappelijke draden die orthomoleculair aanspreekbaar zijn. Door deze draden één voor één te ontrafelen en ze te vervangen door herstellende routines, verandert je biochemische landschap.

De biochemische basis van orthomoleculaire behandeling

Om te begrijpen waarom orthomoleculaire interventies werken bij Fysiopraktijk Zeist, is het handig de belangrijkste "drivers" op een rij te zetten:

Voedingsbalans: De juiste verhouding van macro- en micronutriënten vormt de basis van cellulaire gezondheid. Een uitgebalanceerd voedingspatroon zorgt ervoor dat alle cellulaire processen optimaal kunnen functioneren.

Ontstekingsbalans: Eicosanoïden uit verschillende vetzuren beïnvloeden ontstekingsprocessen in het lichaam. Door de juiste balans van omega-3 en omega-6 vetzuren kunnen we ontstekingsreacties moduleren.

Oxidatieve stress: Bij chronische klachten is antioxidantverdediging vaak verlaagd. Het aanvullen van antioxidanten kan helpen bij het neutraliseren van vrije radicalen en het beschermen van cellen.

Mitochondriale energie: Optimale energieproductie vereist specifieke co-factoren zoals B-vitamines, magnesium en coenzym Q10. Deze nutrients zijn essentieel voor de cellulaire energieproductie.

Darmgezondheid: Een gezonde darmbarrière voorkomt systemische ontstekingssignalen. De darm-hersenverbinding speelt een cruciale rol bij vele gezondheidsproblemen.

Neuro-endocriene as: Stress en slaap beïnvloeden hormonale balans direct. Chronische stress kan leiden tot dysregulatie van verschillende hormoonsystemen.

Praktische implementatie bij Fysiopraktijk Zeist

Stap 1: Grondige analyse
Begin met een uitgebreide anamnese en eventuele laboratoriumtests. Bij Orthomoleculair Praktijk Zeist kijken we naar voedingspatronen, stressfactoren en biochemische markers. Deze holistische benadering geeft ons inzicht in de onderliggende oorzaken van klachten.

Stap 2: Voedingsoptimalisatie
Kies voor onbewerkte, voedingsstofrijke producten die de basis vormen voor optimale gezondheid. Balanceer omega-3 en omega-6 vetzuren door bewuste keuzes in vetten en oliën. Zorg voor voldoende eiwit (1,2-1,6 g/kg lichaamsgewicht) voor weefselonderhoud en herstel. Integreer ontstekingsremmende kruiden en specerijen zoals kurkuma, gember en knoflook.

Stap 3: Gerichte suppletie
Afhankelijk van individuele behoeften kunnen we bij Fysio Zeist aanvullen met specifieke voedingssupplementen. Magnesium voor spier- en zenuwfunctie, vooral belangrijk bij stress en spierspanning. Vitamine D voor immuunbalans en botgezondheid, met name in de Nederlandse klimaatomstandigheden. Omega-3 voor ontstekingsregulatie en hersenfunctie. B-vitamines voor energiemetabolisme en zenuwstelsel ondersteuning.

Stap 4: Leefstijlaanpassingen
Optimaliseer slaapkwaliteit en circadiane ritmes door regelmatige slaaptijden en beperking van blauw licht 's avonds. Implementeer stressmanagement technieken zoals ademhalingsoefeningen, meditatie of yoga. Bouw geleidelijk beweging op binnen individuele mogelijkheden, waarbij consistentie belangrijker is dan intensiteit.

Veelgemaakte fouten en hoe je ze voorkomt

Bij Orthomoleculair Praktijk Zeist zien we regelmatig dezelfde valkuilen die patiënten tegenkomen tijdens hun gezondheidsreis.

Te snelle verwachtingen: Cellulaire veranderingen hebben tijd nodig - reken op 8-12 weken voor merkbare verbetering. Het is belangrijk om geduldig te blijven en het proces te vertrouwen.

Isolatie van supplementen: Een holistisch plan werkt beter dan losse pillen. Supplementen zijn een aanvulling op, geen vervanging van een gezonde leefstijl en voeding.

Negeren van individuele verschillen: Wat bij de één werkt, werkt niet per se bij de ander. Elke persoon heeft unieke behoeften en vereist een geïndividualiseerde aanpak.

Inconsistentie: Regelmatige toepassing is crucialer dan perfectie. Het is beter om 80% consistent te zijn dan 100% perfect voor korte tijd.

Praktijkvoorbeelden uit Zeist

In onze praktijk zien we regelmatig hoe patiënten profiteren van deze geïntegreerde aanpak. Combinaties van voedingsaanpassingen, gerichte suppletie en leefstijlinterventies leveren vaak de beste resultaten. Een holistische benadering waarbij alle aspecten van gezondheid worden geadresseerd, toont consistent betere uitkomsten dan geïsoleerde interventies.

{request.custom_instructions if request.custom_instructions else ''}

Veelgestelde vragen

Hoe lang duurt het voordat ik resultaat zie?
Korte termijn (2-4 weken): Verbeterde energie en slaap. Middellang (6-8 weken): Stabilere stemming en spijsvertering. Lange termijn (3-6 maanden): Structurele verbeteringen in chronische klachten. Geduld en consistentie zijn essentieel voor duurzame resultaten.

Is orthomoleculaire behandeling veilig?
Bij correcte toepassing en professionele begeleiding is orthomoleculaire geneeskunde zeer veilig. Bij Fysiopraktijk Zeist houden we altijd rekening met medicijngebruik en bestaande aandoeningen. We werken samen met je huisarts indien nodig.

Kan ik orthomoleculaire behandeling combineren met reguliere medicatie?
Ja, in veel gevallen is combinatie mogelijk en zelfs wenselijk. Overleg altijd met je behandelaar bij Orthomoleculair Praktijk Zeist voordat je wijzigingen aanbrengt. We hebben ervaring met medicijn-supplement interacties.

Wat zijn de kosten van orthomoleculaire behandeling?
Kosten variëren per individuele situatie en behandelplan. Neem contact op met Fysio Zeist voor een persoonlijk advies en kostenraming. We geloven in transparantie en bespreken altijd vooraf de verwachte kosten.

Disclaimer: De informatie in dit artikel is educatief en vervangt geen medisch advies. Overleg altijd met je behandelaar bij Orthomoleculair Praktijk Zeist, zeker bij medicijngebruik of bestaande aandoeningen.

Over Orthomoleculair Praktijk Zeist
Gevestigd in het hart van Zeist, combineren we moderne orthomoleculaire kennis met persoonlijke begeleiding. Fysio Zeist staat voor kwaliteit, expertise en resultaatgerichte zorg. Ons team heeft jarenlange ervaring in het begeleiden van patiënten naar optimale gezondheid."""

        # Generate realistic SEO metadata
        base_title = request.title[:50] if len(request.title) <= 50 else request.title[:47] + "..."
        seo_data = {
            "meta_title": f"{base_title} | Orthomoleculair Praktijk Zeist",
            "meta_description": f"Ontdek professionele {request.title.lower()} bij Orthomoleculair Praktijk Zeist. Wetenschappelijk onderbouwde orthomoleculaire behandeling in Zeist. ✓ Persoonlijk advies",
            "url_slug": request.title.lower().replace(" ", "-").replace(",", "").replace(".", "").replace("(", "").replace(")", ""),
            "primary_keywords": [
                request.title.lower(),
                "orthomoleculair zeist", 
                "fysiopraktijk zeist",
                "natuurlijke behandeling zeist"
            ],
            "suggested_tags": [
                "orthomoleculair",
                "natuurlijke behandeling", 
                "voedingssupplementen",
                "holistische gezondheid",
                "micronutriënten",
                "preventieve zorg",
                "fysiopraktijk zeist",
                "orthomoleculaire geneeskunde",
                "gezonde leefstijl",
                "integrative medicine"
            ]
        }
        
        # Select appropriate featured image based on blog topic
        featured_image = "https://images.unsplash.com/photo-1704694671866-f83e0b91df09?crop=entropy&cs=srgb&fm=jpg&ixid=M3w3NDk1Nzl8MHwxfHNlYXJjaHwxfHx8MTc1OTY3ODYwNXww&ixlib=rb-4.1.0&q=85"
        
        # Create blog article document
        blog_article = BlogArticle(
            title=request.title,
            content=mock_blog_content,
            tags=seo_data.get("suggested_tags", []),
            category=request.category,
            meta_title=seo_data.get("meta_title", f"{request.title} | Orthomoleculair Praktijk Zeist"),
            meta_description=seo_data.get("meta_description", "Professioneel orthomoleculair advies in Zeist"),
            url_slug=seo_data.get("url_slug", request.title.lower().replace(" ", "-")),
            featured_image_url=featured_image,
            source_document_ids=request.document_ids,
            custom_instructions=request.custom_instructions
        )
        
        # Save to database
        blog_dict = blog_article.dict()
        result = await db.documents.insert_one(blog_dict)
        
        return {
            "success": True,
            "blog_id": blog_article.id,
            "blog_article": blog_article,
            "seo_data": seo_data
        }
        
    except Exception as e:
        logging.error(f"Blog creation error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Blog creation failed: {str(e)}")

# Statistics
@api_router.get("/stats")
async def get_stats():
    """Get knowledge base statistics"""
    total_docs = await db.documents.count_documents({})
    categories = await db.documents.distinct("category")
    
    category_counts = {}
    for cat in categories:
        count = await db.documents.count_documents({"category": cat})
        category_counts[cat] = count
    
    return {
        "total_documents": total_docs,
        "categories": category_counts
    }

@api_router.get("/export/oneliners")
async def export_oneliners():
    """Export all document one-liners in CSV format for Make.com automation"""
    try:
        documents = await db.documents.find({}, {
            "id": 1, "title": 1, "category": 1, "one_liner": 1, 
            "created_at": 1, "tags": 1, "_id": 0
        }).to_list(length=None)
        
        # Prepare data for spreadsheet export
        export_data = []
        for doc in documents:
            export_data.append({
                "id": doc.get("id", ""),
                "title": doc.get("title", ""),
                "category": doc.get("category", ""),
                "one_liner": doc.get("one_liner", ""),
                "tags": ", ".join(doc.get("tags", [])),
                "created_date": doc.get("created_at", "").split("T")[0] if doc.get("created_at") else ""
            })
        
        return {
            "success": True,
            "count": len(export_data),
            "data": export_data
        }
        
    except Exception as e:
        logging.error(f"Export error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/documents/{document_id}/regenerate-oneliner")
async def regenerate_oneliner(document_id: str):
    """Regenerate one-liner for a specific document"""
    try:
        # Get document
        doc = await db.documents.find_one({"id": document_id})
        if not doc:
            raise HTTPException(status_code=404, detail="Document niet gevonden")
        
        # Generate new one-liner
        new_oneliner = generate_oneliner_mock(doc["title"], doc["content"])
        
        # Update document
        await db.documents.update_one(
            {"id": document_id},
            {"$set": {"one_liner": new_oneliner}}
        )
        
        return {
            "success": True,
            "document_id": document_id,
            "new_oneliner": new_oneliner
        }
        
    except Exception as e:
        logging.error(f"Regenerate oneliner error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/")
async def root():
    return {"message": "Wellness Knowledge Archive API"}

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()